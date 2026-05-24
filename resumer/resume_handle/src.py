import os
from enum import Enum
from docx import Document


class Section(Enum):
    PROJECTS = 1
    EDUCATION = 2
    WORK_EXPERIENCE = 3
    TECHNICAL_SKILLS = 4
    LEADERSHIP_AND_INVOLVEMENT = 5


class ResumeHandler:
    def __init__(self, loc: str) -> None:
        if not os.path.exists(loc):
            raise Exception("Could not find resume")
        if os.path.splitext(loc)[-1].lower() != ".docx":
            raise Exception("File is not a DOCX file")

        self.loc = loc
        self.doc = Document(self.loc)

    def get_run_from_text(self, text) -> tuple[int, int]:
        """
        Returns coordinates to the found run,

        tuple(paragraph index, run index relative to the paragraph index)
        """
        p_idx = 0

        for p in self.doc.paragraphs:

            r_idx = 0

            for run in p.runs:

                if run.text == text:
                    return p_idx, r_idx

                r_idx += 1
            p_idx += 1

        return -1, -1

    def get_next_run(self, p_idx, r_idx):
        """
        Returns the coordinates of the
        next valid run, regardless if on
        another paragraph. Otherwise, returns
        (-1, -1).
        """
        if (r_idx + 1) < len(self.doc.paragraphs[p_idx].runs):
            return p_idx, r_idx + 1

        for next_p in range(p_idx + 1, len(self.doc.paragraphs)):
            if self.doc.paragraphs[next_p].runs:
                return next_p, 0

        return -1, -1

    def get_paragraph_with_text(self, text) -> int:
        """
        Returns the index to the found paragraph
        """
        idx = 0

        for para in self.doc.paragraphs:
            if text in para.text:
                return idx
            idx += 1

        return idx
